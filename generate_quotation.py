import sys
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_quotation(data, output_path):
    doc = Document()

    # Add Document Number and Date
    top_table = doc.add_table(rows=1, cols=2)
    top_table.cell(0, 0).text = f"Document: {data['document_number']}"
    top_table.cell(0, 1).text = f"Date: {data['date']}"
    top_table.cell(0, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph('')  # Spacing

    # Add Recipient and Product Information
    doc.add_paragraph(f"To: {data['to']}")
    doc.add_paragraph(f"Subject: {data['subject']}")
    doc.add_paragraph(f"Reference: {data['reference']}")
    doc.add_paragraph('')

    table = doc.add_table(rows=2, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Total Price'

    row_cells = table.rows[1].cells
    row_cells[0].text = data['product_description']
    row_cells[1].text = str(data['quantity'])
    row_cells[2].text = f"{data['unit_price']:.2f}"
    row_cells[3].text = f"{data['quantity'] * data['unit_price']:.2f}"

    doc.save(output_path)
    print(f"Quotation saved at {output_path}")

if __name__ == "__main__":
    json_data = json.loads(sys.argv[1])
    output_path = sys.argv[2]
    generate_quotation(json_data, output_path)
